from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

name = 'Eduard Haponov'
position = 'RPA Developer'
location = 'Wexford, County Wexford, Ireland'
email = 'fvtvlnight@gmail.com'
phone = '087 392 7134'
github = 'Github.com/Fvtvl'
linkedin = 'Linkedin.com/in/fvtvlnight'
skills = [
    'Strong knowledge of business context and processes.',
    'Experience in designing, developing, testing, and deploying Robotics solutions with UiPath and BluePrism',
    'Robotics configuration - architecture, design, development, and maintenance.',
    'Programming Languages: JavaScript, HTML, CSS.',
    'JavaScript libraries: Node.js, React.js.',
    'Database Systems: Firebase, MongoDB.',
    'Control version: Git.',
    'APIs: REST.',
    'Applications: UiPath, BluePrism'
]
experience = [
    ('1 year', 'RPA Developer'),
    ('1 year', 'Frontend Developer '),
    ('3 years', 'Private entrepreneur')
]


def create_cv(name, position, location, email, phone, github, linkedin, skills, experience):
    doc = Document()

    title = doc.add_heading(name, level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(f'{position}')
    doc.add_paragraph(f'{location}')
    doc.add_paragraph(f'Email: {email}')
    doc.add_paragraph(f'Phone: {phone}')
    doc.add_paragraph(f'Github: {github}')
    doc.add_paragraph(f'Linkedin: {linkedin}')

    doc.add_heading('Skills', level=2)
    skills_paragraph = doc.add_paragraph()
    skills_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for skill in skills:
        skills_paragraph.add_run(f'- {skill}\n')

    doc.add_heading('Experience', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Years'
    hdr_cells[1].text = 'Title'
    for title, description in experience:
        row_cells = table.add_row().cells
        row_cells[0].text = title
        row_cells[1].text = description

    return doc

cv = create_cv(name, position, location, email, phone, github, linkedin, skills, experience)

cv.save('cv.docx')
